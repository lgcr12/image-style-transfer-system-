from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont


BASE_DIR = Path(r"E:\models")
DOC_PATH = BASE_DIR / "项目总体设计图与数据库表说明.docx"
DIAGRAM_PATH = BASE_DIR / "项目总体设计图.png"


def load_font(size: int):
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for font_path in candidates:
        p = Path(font_path)
        if p.exists():
            try:
                return ImageFont.truetype(str(p), size)
            except Exception:
                continue
    return ImageFont.load_default()


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=6, align="center")
    text_w = bbox[2] - bbox[0]
    text_h = bbox[3] - bbox[1]
    x = x1 + (x2 - x1 - text_w) / 2
    y = y1 + (y2 - y1 - text_h) / 2
    draw.multiline_text((x, y), text, font=font, fill=fill, spacing=6, align="center")


def make_diagram() -> None:
    width, height = 1800, 1080
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    font_title = load_font(42)
    font_box = load_font(28)
    font_small = load_font(22)

    blue = (53, 99, 233)
    light = (240, 245, 255)
    dark = (40, 48, 61)
    gray = (109, 120, 135)
    line = (125, 145, 175)

    def box(x1: int, y1: int, x2: int, y2: int, title: str, subtitle: str = ""):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=light, outline=blue, width=3)
        draw_centered_text(draw, (x1 + 10, y1 + 18, x2 - 10, y1 + 68), title, font_box, dark)
        if subtitle:
            draw_centered_text(draw, (x1 + 15, y1 + 70, x2 - 15, y2 - 15), subtitle, font_small, gray)

    def arrow(x1: int, y1: int, x2: int, y2: int, label: str = ""):
        draw.line((x1, y1, x2, y2), fill=line, width=5)
        import math

        angle = math.atan2(y2 - y1, x2 - x1)
        arrow_head = 16
        points = [
            (x2, y2),
            (
                x2 - arrow_head * math.cos(angle - 0.45),
                y2 - arrow_head * math.sin(angle - 0.45),
            ),
            (
                x2 - arrow_head * math.cos(angle + 0.45),
                y2 - arrow_head * math.sin(angle + 0.45),
            ),
        ]
        draw.polygon(points, fill=line)
        if label:
            mid_x = (x1 + x2) / 2
            mid_y = (y1 + y2) / 2
            draw.text((mid_x - 28, mid_y - 26), label, font=font_small, fill=gray)

    draw.text((60, 30), "项目总体设计图", fill=dark, font=font_title)

    box(80, 180, 320, 300, "用户浏览器", "上传图片\n查看结果")
    box(420, 160, 760, 320, "前端页面", "templates + static")
    box(860, 140, 1210, 340, "FastAPI 服务", "app.py")
    box(1300, 120, 1670, 280, "任务队列", "JobQueue\n单任务串行处理")
    box(1300, 340, 1670, 500, "推理模块", "style_transfer.py\nsd_style_transfer.py")
    box(860, 420, 1210, 580, "分析与评分", "image_analyzer.py\nrecipe_scorer.py")
    box(420, 430, 760, 590, "导出与分享", "exporter.py\nshare_builder.py")
    box(860, 690, 1210, 860, "SQLite 数据库", "data/jobs.db")
    box(420, 700, 760, 860, "文件存储", "uploads / results / meta")
    box(80, 700, 320, 860, "说明", "数据库保存索引信息\n图片结果保存在磁盘")

    arrow(320, 240, 420, 240)
    arrow(760, 240, 860, 240)
    arrow(1210, 220, 1300, 200, "入队")
    arrow(1485, 280, 1485, 340, "执行")
    arrow(1300, 420, 1210, 500)
    arrow(860, 500, 760, 510)
    arrow(1035, 580, 1035, 690)
    arrow(860, 270, 760, 760)
    arrow(1035, 340, 1035, 690)

    image.save(DIAGRAM_PATH)


def set_chinese_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_paragraph(document: Document, text: str, first_line: bool = True):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.line_spacing = 1.5
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    return paragraph


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    run.bold = bold


def add_table(document: Document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    document.add_paragraph("")


def build_doc() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)
    set_chinese_style(document)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run1 = title.add_run("项目总体设计图与数据库表说明\n")
    run1.bold = True
    run1.font.size = Pt(18)
    run1.font.name = "黑体"
    run1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run2 = title.add_run("课程作业精简版")
    run2.font.size = Pt(14)
    run2.font.name = "黑体"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    document.add_heading("一、总体设计图", level=1)
    add_paragraph(
        document,
        "本项目总体上采用“前端页面、后端接口、任务调度、模型推理、数据存储”五部分组成的结构。用户通过浏览器上传图片并选择参数后，请求会发送到 FastAPI 服务端。服务端负责接收数据、创建任务，并根据不同的处理方式调用对应的推理模块。任务执行完成以后，结果图片保存在本地目录中，任务信息和索引数据则保存在 SQLite 数据库中。",
    )
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    picture_paragraph.add_run().add_picture(str(DIAGRAM_PATH), width=Inches(6.5))
    caption = document.add_paragraph("图 1  项目总体设计图")
    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_paragraph(
        document,
        "从图 1 可以看出，浏览器和前端页面属于用户交互层，FastAPI 服务属于系统控制层，任务队列和推理模块属于处理层，SQLite 数据库和文件目录属于存储层。这样的划分使系统结构比较清楚，也方便后续继续增加新的模型和页面功能。",
    )
    for sentence in [
        "前端部分由 templates 和 static 目录组成，主要负责页面展示、参数输入和结果显示。",
        "后端部分集中在 app.py 中，负责接口路由、任务创建和结果返回。",
        "任务队列模块用于管理高开销任务，避免多个重任务同时运行导致系统不稳定。",
        "推理模块包含传统风格迁移和 SD 风格迁移两种处理方式，是项目的核心业务部分。",
        "数据存储采用数据库和文件目录结合的方式，数据库保存任务信息，文件目录保存图片结果。",
    ]:
        add_paragraph(document, sentence, first_line=False)

    document.add_heading("二、数据库表设计", level=1)
    add_paragraph(
        document,
        "本项目使用 SQLite 作为数据库，数据库文件位于 data/jobs.db。经过检查，当前数据库中主要有两张业务表，分别是 jobs 表和 gallery_items 表。整体设计比较简单，适合本地项目和课程作业场景。",
    )
    document.add_heading("1. 数据库关系说明", level=2)
    add_paragraph(
        document,
        "jobs 表是系统中的核心表，用来记录每一次任务的状态、参数和结果信息；gallery_items 表则用于保存已经发布到画廊中的展示信息。两张表通过 job_id 字段建立联系，因此可以理解为一条任务记录最多对应一条画廊记录。",
    )
    relation = document.add_paragraph()
    relation.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    relation_run = relation.add_run("JOBS  1  ------  0..1  GALLERY_ITEMS")
    relation_run.font.name = "Consolas"
    relation_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    relation_run.font.size = Pt(11)
    relation_caption = document.add_paragraph("图 2  数据表关系示意")
    relation_caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_heading("2. jobs 表设计", level=2)
    add_paragraph(
        document,
        "jobs 表相当于系统的任务主表，无论是传统风格迁移任务还是 SD 风格迁移任务，最终都会写入这张表。该表不仅记录任务状态，还保存用户提交参数、结果图片路径以及评分信息。",
    )
    add_table(
        document,
        ["字段名", "类型", "说明"],
        [
            ["job_id", "TEXT", "任务唯一编号，主键"],
            ["mode", "TEXT", "任务类型，如 style-transfer 或 sd"],
            ["status", "TEXT", "任务状态，如 queued、running、finished、error"],
            ["phase", "TEXT", "任务当前阶段"],
            ["phase_detail", "TEXT", "阶段详细说明"],
            ["progress", "INTEGER", "任务进度"],
            ["error", "TEXT", "错误信息"],
            ["score_json", "TEXT", "结果评分，采用 JSON 文本保存"],
            ["next_recipe_json", "TEXT", "推荐参数信息，采用 JSON 文本保存"],
            ["params_json", "TEXT", "用户提交参数，采用 JSON 文本保存"],
            ["content_path", "TEXT", "原始内容图路径"],
            ["style_path", "TEXT", "风格图路径，可为空"],
            ["result_path", "TEXT", "主结果图路径"],
            ["result_paths_json", "TEXT", "多结果路径列表"],
            ["created_at_ms", "INTEGER", "创建时间"],
            ["updated_at_ms", "INTEGER", "更新时间"],
        ],
    )
    add_paragraph(
        document,
        "从表结构可以看出，jobs 表既保存普通字段，也保存 JSON 文本字段。这种设计的优点是实现简单、扩展方便，能够满足当前项目对任务记录和结果管理的基本需求。",
    )

    document.add_heading("3. gallery_items 表设计", level=2)
    add_paragraph(
        document,
        "gallery_items 表主要用于保存画廊展示信息。它并不重复保存完整任务数据，而是只保留标题、匿名状态和发布时间等展示字段，再通过 job_id 关联 jobs 表中的详细记录。这样可以减少数据重复。",
    )
    add_table(
        document,
        ["字段名", "类型", "说明"],
        [
            ["id", "INTEGER", "自增主键"],
            ["job_id", "TEXT", "对应任务编号，唯一"],
            ["title", "TEXT", "画廊展示标题"],
            ["anonymous", "INTEGER", "是否匿名发布，0 表示否，1 表示是"],
            ["created_at_ms", "INTEGER", "发布时间"],
        ],
    )
    add_paragraph(
        document,
        "需要注意的是，gallery_items.job_id 在业务上依赖 jobs.job_id，但当前数据库中没有设置真正的外键约束，因此它属于一种较为简单的软关联设计。",
    )

    document.add_heading("4. 索引设计", level=2)
    add_paragraph(
        document,
        "为了提高查询效率，项目还建立了几个常用索引，主要用于支持历史记录查询、任务筛选和画廊列表展示。",
    )
    add_table(
        document,
        ["索引名", "所在表", "作用说明"],
        [
            ["idx_jobs_created", "jobs", "按创建时间倒序查询任务记录"],
            ["idx_jobs_mode_status", "jobs", "按任务类型和状态筛选任务"],
            ["idx_gallery_created", "gallery_items", "按发布时间倒序查询画廊内容"],
        ],
    )

    document.save(DOC_PATH)


if __name__ == "__main__":
    make_diagram()
    build_doc()
    print(DOC_PATH)
